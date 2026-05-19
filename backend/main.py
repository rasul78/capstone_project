"""
Sentinel AI — FastAPI Backend v4.0 (Multi-Agent Edition)

Routes:
  GET  /                          → health check
  GET  /api/model/info            → информация о VisionNet
  POST /api/train/start           → запуск обучения
  POST /api/train/stop            → остановка
  GET  /api/train/status          → статус + история
  GET  /api/train/stream          → SSE прогресс
  GET  /api/train/history         → история запусков из PostgreSQL
  POST /api/predict               → классификация изображения
  GET  /api/kb/stats              → статистика базы знаний
  GET  /api/kb/documents          → список документов
  POST /api/kb/documents          → добавить документ (JSON)
  POST /api/kb/documents/upload   → загрузить файл
  DELETE /api/kb/documents/{id}   → удалить документ
  POST /api/kb/chat               → вопрос к базе знаний (legacy)
  GET  /api/kb/chat/history       → история чатов из PostgreSQL
  POST /api/agent/chat            → вопрос через оркестратор (RAG + Web + Synthesis)
  GET  /api/agent/metrics         → observability метрики
  GET  /api/agent/traces          → последние трейсы запросов
  POST /api/kb/search             → семантический поиск
"""

import asyncio, io, json, base64, threading, logging, urllib.request, urllib.error
from typing import Optional

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from PIL import Image

import sys, os
sys.path.insert(0, os.path.dirname(__file__))

# ── Load .env ──────────────────────────────────────────────────────────────
def _load_env():
    paths = [
        os.path.join(os.path.dirname(__file__), '.env'),
        os.path.join(os.path.dirname(__file__), '..', '.env'),
        '.env',
    ]
    for p in paths:
        if os.path.exists(p):
            with open(p, encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if line and not line.startswith('#') and '=' in line:
                        k, v = line.split('=', 1)
                        k = k.strip()
                        v = v.strip().strip('"').strip("'")
                        os.environ[k] = v  # всегда перезаписываем
            # Показываем какие ключи загружены
            groq = os.getenv('GROQ_API_KEY', '')
            ant  = os.getenv('ANTHROPIC_API_KEY', '')
            gem  = os.getenv('GEMINI_API_KEY', '')
            print(f"[ENV] Загружен: {p}")
            print(f"[ENV] GROQ_API_KEY:      {'✅ ' + groq[:12] + '...' if groq else '❌ не задан'}")
            print(f"[ENV] ANTHROPIC_API_KEY: {'✅ задан' if ant else '❌ не задан'}")
            print(f"[ENV] GEMINI_API_KEY:    {'✅ задан' if gem else '❌ не задан'}")
            return
    print("[ENV] ❌ .env файл не найден!")
    print("[ENV] Создайте файл backend/.env с содержимым:")
    print("[ENV]   GROQ_API_KEY=gsk_ваш_ключ")
_load_env()

from models.vision import VisionTrainer
from models.knowledge import KnowledgeBase
from database import (
    init_db, close_pool,
    db_save_chat, db_get_chat_history,
    db_create_training_run, db_finish_training_run, db_get_training_runs,
)
from agents.orchestrator import AgentOrchestrator
from agents.langgraph_orchestrator import LangGraphOrchestrator
from streaming_chat import add_streaming_routes
from mcp_server import SentinelMCPServer, create_mcp_router

# ── НОВОЕ v4.1: auth + integrations ───────────────────────────────────
from auth_routes import build_auth_router
from auth.auth import ensure_demo_users
from chat_fast_v2 import build_chat_router
from integrations.cache import init_cache, close_cache, cache_stats
from integrations.reranker import reranker_status
from integrations.train_ws import train_ws, train_ws_endpoint
from integrations.stream_chat import build_stream_router as _old_stream_router_unused

# v4.3: Streaming с прогрессом агентов (Variant 9)
try:
    from stream_chat_v2 import build_stream_router as build_agentic_stream_router
    _AGENTIC_STREAM_OK = True
except ImportError as e:
    _AGENTIC_STREAM_OK = False
    logger.warning("Agentic stream not available: %s", e)
from integrations.telegram_notify import notify as tg_notify, is_configured as tg_on

# ── MCP integration (v4.2) ────────────────────────────────────────────
try:
    from mcp_routes import init_mcp_registry, close_mcp_registry, build_mcp_status_router
    _MCP_INTEGRATION_OK = True
except ImportError as e:
    _MCP_INTEGRATION_OK = False
    logger.warning("MCP integration not available: %s", e)

# ── Logging ───────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s — %(message)s",
)
logger = logging.getLogger("sentinel.main")

# ── App ───────────────────────────────────────────────────────────────────
app = FastAPI(title="Sentinel AI", version="4.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Singletons ────────────────────────────────────────────────────────────
_trainer:      Optional[VisionTrainer]     = None
_train_log:    list                        = []
_train_thread: Optional[threading.Thread] = None
_run_id:       Optional[int]              = None
kb           = KnowledgeBase(device='cpu')
_default_trainer: VisionTrainer = None
orchestrator: Optional[AgentOrchestrator] = None


@app.on_event("startup")
async def startup():
    global orchestrator
    try:
        await asyncio.wait_for(init_db(), timeout=10.0)
    except asyncio.TimeoutError:
        logger.warning("⚠️  DB init timeout — используем SQLite fallback")
    except Exception as e:
        logger.error("⚠️  DB init error: %s — продолжаем без БД", e)

    # ── НОВОЕ: создаём демо-юзеров (demo/demo123, admin/admin123) ─────
    try:
        await ensure_demo_users()
    except Exception as e:
        logger.warning("⚠️  ensure_demo_users error: %s", e)

    # ── НОВОЕ: инициализируем кэш (Redis или in-memory) ───────────────
    try:
        await init_cache()
    except Exception as e:
        logger.warning("⚠️  init_cache error: %s", e)

    try:
        await asyncio.wait_for(kb.load_from_db(), timeout=10.0)
    except Exception as e:
        logger.warning("⚠️  KB load error: %s", e)

    orchestrator = LangGraphOrchestrator(kb)

    mcp_router = create_mcp_router(orchestrator.mcp_server)
    app.include_router(mcp_router)

    # ── НОВОЕ: подключаем auth + chat v2 + stream + ws ────────────────
    app.include_router(build_auth_router())
    app.include_router(build_chat_router(kb, orchestrator))
    # v4.3: Agentic streaming with per-agent progress events
    if _AGENTIC_STREAM_OK:
        app.include_router(build_agentic_stream_router(kb, orchestrator))
    app.add_api_websocket_route("/ws/train", train_ws_endpoint)

    # ── НОВОЕ (v4.2): MCP registry + status endpoints ──────────────
    if _MCP_INTEGRATION_OK:
        await init_mcp_registry()
        app.include_router(build_mcp_status_router())

    logger.info("✅ Sentinel AI запущен с LangGraph + MCP + Auth + Cache + Reranker")
    logger.info("   Граф: rewrite_query → search_kb → [search_web] → synthesize")
    logger.info("   MCP инструментов: %d", len(orchestrator.mcp_server.list_tools()))
    logger.info("   Ollama: %s", orchestrator.synthesis._model or "недоступна")
    logger.info("   Cache: %s", cache_stats())
    logger.info("   Reranker: %s", reranker_status())
    logger.info("   Telegram: %s", "ON" if tg_on() else "OFF")

    if tg_on():
        try:
            await tg_notify("🚀 <b>Sentinel AI</b> запущен", level="success", silent=True)
        except Exception:
            pass


@app.on_event("shutdown")
async def shutdown():
    await close_pool()
    await close_cache()
    if _MCP_INTEGRATION_OK:
        await close_mcp_registry()


def get_trainer(dataset: str = 'imagenet') -> VisionTrainer:
    global _trainer
    if _trainer is None or _trainer.dataset != dataset:
        _trainer = VisionTrainer(dataset=dataset, device='cpu', save_dir='./checkpoints')
        _train_log.clear()
    return _trainer


# ── Pydantic schemas ──────────────────────────────────────────────────────
class TrainRequest(BaseModel):
    dataset:      str   = Field(default='cifar10')
    epochs:       int   = Field(default=15, ge=1, le=100)
    lr:           float = Field(default=0.001, gt=0, le=1.0)
    batch_size:   int   = Field(default=64, ge=8, le=256)
    weight_decay: float = Field(default=1e-4, ge=0)

class ChatRequest(BaseModel):
    question:   str = Field(min_length=1, max_length=500)
    session_id: str = Field(default='default')

class AgentChatRequest(BaseModel):
    question:   str            = Field(min_length=1, max_length=500)
    session_id: str            = Field(default='default')
    feedback:   Optional[int] = Field(default=None, ge=1, le=5)

class AddDocRequest(BaseModel):
    name:     str = Field(min_length=1, max_length=200)
    content:  str = Field(min_length=10)
    category: str = Field(default='Общее')


# ══════════════════════════════════════════════════════════════════════════
# ── HEALTH ────────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════

@app.get("/", tags=["health"])
def root():
    return {
        "status":  "ok",
        "service": "Sentinel AI",
        "version": "4.0.0",
        "modules": ["VisionNet CNN", "Knowledge Base RAG", "PostgreSQL",
                    "ResearchAgent", "WebAgent", "SynthesisAgent"],
    }


# ══════════════════════════════════════════════════════════════════════════
# ── VISION CNN ────────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════

@app.get("/api/model/info", tags=["vision"])
def model_info():
    return get_trainer().info()


@app.post("/api/train/start", tags=["vision"])
async def train_start(req: TrainRequest):
    global _trainer, _train_log, _train_thread, _run_id

    if _trainer and _trainer.training:
        raise HTTPException(400, "Обучение уже запущено.")

    _trainer = VisionTrainer(dataset=req.dataset if req.dataset else 'imagenet', device='cpu', save_dir='./checkpoints')
    _train_log.clear()

    # ── НОВОЕ: callback пишет и в лог, и в WebSocket ──────────────────
    def _train_cb(d):
        _train_log.append(d)
        try:
            train_ws.broadcast_nowait({"type": "progress", **d})
        except Exception:
            pass
    _trainer.callback = _train_cb

    _run_id = await db_create_training_run(req.dataset, req.epochs, req.lr, req.batch_size)

    def run():
        _trainer.train(
            epochs=req.epochs, lr=req.lr,
            batch_size=req.batch_size, weight_decay=req.weight_decay,
        )
        asyncio.run(db_finish_training_run(
            _run_id,
            best_acc=_trainer.info().get('best_val_acc', 0),
            history=_trainer.history,
        ))
        # WebSocket: финал
        try:
            train_ws.broadcast_nowait({
                "type": "finished",
                "best_acc": _trainer.info().get('best_val_acc', 0),
                "run_id": _run_id,
            })
        except Exception:
            pass
        # Telegram уведомление
        try:
            if tg_on():
                acc = _trainer.info().get('best_val_acc', 0)
                asyncio.run(tg_notify(
                    f"🎯 Обучение завершено\n"
                    f"Dataset: <b>{req.dataset}</b>\n"
                    f"Best accuracy: <b>{acc:.2%}</b>",
                    level="success",
                ))
        except Exception:
            pass

    _train_thread = threading.Thread(target=run, daemon=True)
    _train_thread.start()

    return {"status": "started", "dataset": req.dataset,
            "epochs": req.epochs, "lr": req.lr, "run_id": _run_id}


@app.post("/api/train/stop", tags=["vision"])
def train_stop():
    if _trainer:
        _trainer.stop()
    return {"status": "stopped"}


@app.get("/api/train/status", tags=["vision"])
def train_status():
    if not _trainer:
        return {"is_training": False, "epoch": 0, "best_val_acc": 0, "history": {}, "log": []}
    info = _trainer.info()
    return {
        "is_training":  _trainer.training,
        "epoch":        _trainer.epoch,
        "best_val_acc": info["best_val_acc"],
        "history":      _trainer.history,
        "log":          _train_log[-30:],
    }


@app.get("/api/train/stream", tags=["vision"])
async def train_stream():
    async def event_gen():
        sent = 0
        while True:
            if len(_train_log) > sent:
                for entry in _train_log[sent:]:
                    yield f"data: {json.dumps(entry, ensure_ascii=False)}\n\n"
                sent = len(_train_log)
            if _trainer and not _trainer.training and sent > 0:
                yield f"data: {json.dumps({'done': True})}\n\n"
                break
            await asyncio.sleep(0.4)
    return StreamingResponse(
        event_gen(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.get("/api/train/history", tags=["vision"])
async def train_history():
    runs = await db_get_training_runs()
    return {"runs": [dict(r) for r in runs]}


@app.post("/api/predict", tags=["vision"])
async def predict(file: UploadFile = File(...)):
    raw = await file.read()
    try:
        img = Image.open(io.BytesIO(raw))
    except Exception:
        raise HTTPException(400, "Не удалось открыть изображение.")
    trainer = get_trainer()
    result  = trainer.predict(img)
    preview = img.convert('RGB').resize((128, 128))
    buf     = io.BytesIO()
    preview.save(buf, format='JPEG', quality=80)
    result['preview']  = 'data:image/jpeg;base64,' + base64.b64encode(buf.getvalue()).decode()
    result['filename'] = file.filename or 'image'
    return result


# ══════════════════════════════════════════════════════════════════════════
# ── KNOWLEDGE BASE ────────────────────────────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════

@app.get("/api/kb/stats", tags=["knowledge"])
def kb_stats():
    return kb.stats()


@app.get("/api/kb/documents", tags=["knowledge"])
def kb_list():
    return {"documents": kb.list_documents()}


@app.post("/api/kb/documents", tags=["knowledge"])
async def kb_add(req: AddDocRequest):
    doc = await kb.add_document_async(req.name.strip(), req.content.strip(), req.category)
    return {"success": True, "document": doc, "stats": kb.stats()}


@app.post("/api/kb/documents/upload", tags=["knowledge"])
async def kb_upload(file: UploadFile = File(...), category: str = "Загружено"):
    raw      = await file.read()
    filename = file.filename or "document"
    ext      = filename.rsplit(".", 1)[-1].lower()
    text     = None

    # ── PDF ──────────────────────────────────────────────────────────────
    if ext == "pdf":
        try:
            import fitz  # pymupdf
            doc_pdf     = fitz.open(stream=raw, filetype="pdf")
            pages_text  = []
            total_pages = len(doc_pdf)
            print(f"[PDF] Открыт: {filename}, страниц: {total_pages}")

            for page_num, page in enumerate(doc_pdf):
                pt = page.get_text("text")
                print(f"[PDF] Страница {page_num + 1}: {len(pt)} символов")
                if pt.strip():
                    pages_text.append(f"[Страница {page_num + 1}]\n{pt}")

            doc_pdf.close()
            text = "\n\n".join(pages_text)
            print(f"[PDF] pymupdf итого: {len(text)} символов")

            # Fallback: scanned PDF без текстового слоя
            if not text.strip():
                import pdfplumber as _plumber
                with _plumber.open(io.BytesIO(raw)) as pdf_doc:
                    pages_text = []
                    for page_num, page in enumerate(pdf_doc.pages):
                        pt = page.extract_text()
                        if pt and pt.strip():
                            pages_text.append(f"[Страница {page_num + 1}]\n{pt}")
                text = "\n\n".join(pages_text)
                print(f"[PDF] pdfplumber итого: {len(text)} символов")

            # OCR цепочка: Tesseract -> EasyOCR -> Claude Vision
            if not text.strip():
                print("[PDF] Сканированный PDF — применяю OCR...")

                # 1. Tesseract (локальный, бесплатный, лучший для кириллицы)
                try:
                    import pytesseract, shutil
                    from PIL import Image as _PI, ImageEnhance as _PE
                    import fitz as _ft
                    tess = shutil.which("tesseract")
                    if not tess:
                        for _tp in [
                            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                        ]:
                            if os.path.exists(_tp):
                                tess = _tp
                                break
                    if tess:
                        pytesseract.pytesseract.tesseract_cmd = tess
                        _doc = _ft.open(stream=raw, filetype="pdf")
                        _pts = []
                        for _pn, _pg in enumerate(_doc):
                            _pix = _pg.get_pixmap(matrix=_ft.Matrix(3, 3), colorspace=_ft.csRGB)
                            _img = _PI.frombytes("RGB", [_pix.width, _pix.height], _pix.samples)
                            _img = _PE.Contrast(_img).enhance(2.0)
                            _t = pytesseract.image_to_string(_img, lang="rus+eng", config="--psm 6")
                            if _t.strip():
                                _pts.append("[Страница " + str(_pn+1) + "]\n" + _t.strip())
                        _doc.close()
                        text = "\n\n".join(_pts)
                        if text.strip():
                            print("[OCR] Tesseract: " + str(len(text)) + " символов")
                    else:
                        print("[OCR] Tesseract не найден. Запусти install_tesseract.bat")
                except ImportError:
                    print("[OCR] pytesseract не установлен: pip install pytesseract")
                except Exception as _te:
                    print("[OCR] Tesseract error: " + str(_te))

                # 2. EasyOCR
                if not text.strip():
                    try:
                        import easyocr as _eo, numpy as _np, fitz as _fe
                        _r    = _eo.Reader(["ru", "en"], gpu=False, verbose=False)
                        _doc2 = _fe.open(stream=raw, filetype="pdf")
                        _eps  = []
                        for _pn, _pg in enumerate(_doc2):
                            _pix = _pg.get_pixmap(matrix=_fe.Matrix(2, 2))
                            _arr = _np.frombuffer(_pix.samples, dtype=_np.uint8).reshape(_pix.height, _pix.width, _pix.n)
                            _res = _r.readtext(_arr, detail=0, paragraph=True)
                            _t   = "\n".join(_res)
                            if _t.strip():
                                _eps.append("[Страница " + str(_pn+1) + "]\n" + _t)
                        _doc2.close()
                        text = "\n\n".join(_eps)
                        if text.strip():
                            print("[OCR] EasyOCR: " + str(len(text)) + " символов")
                    except ImportError:
                        print("[OCR] easyocr не установлен: pip install easyocr")
                    except Exception as _ee:
                        print("[OCR] EasyOCR error: " + str(_ee))

                # 3. Claude Vision API
                _ak = os.getenv("ANTHROPIC_API_KEY", "").strip().strip('"')
                if _ak and not text.strip():
                    try:
                        import fitz as _fc
                        _doc3 = _fc.open(stream=raw, filetype="pdf")
                        _cps  = []
                        for _pn, _pg in enumerate(_doc3):
                            _pix = _pg.get_pixmap(matrix=_fc.Matrix(200/72, 200/72), colorspace=_fc.csRGB)
                            _b64 = base64.b64encode(_pix.tobytes("jpeg")).decode()
                            _pl  = json.dumps({
                                "model": "claude-haiku-4-5-20251001", "max_tokens": 2048,
                                "messages": [{"role": "user", "content": [
                                    {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg", "data": _b64}},
                                    {"type": "text", "text": "Extract all text from this document page. Return only the text."}
                                ]}]
                            }).encode()
                            _req = urllib.request.Request(
                                "https://api.anthropic.com/v1/messages", data=_pl,
                                headers={"x-api-key": _ak, "anthropic-version": "2023-06-01", "Content-Type": "application/json"},
                                method="POST"
                            )
                            _loop = asyncio.get_event_loop()
                            _resp = await _loop.run_in_executor(None, lambda r=_req: urllib.request.urlopen(r, timeout=30))
                            _t = json.loads(_resp.read())["content"][0]["text"].strip()
                            if _t:
                                _cps.append("[Страница " + str(_pn+1) + "]\n" + _t)
                        _doc3.close()
                        text = "\n\n".join(_cps)
                        if text.strip():
                            print("[OCR] Claude Vision: " + str(len(text)) + " символов")
                    except Exception as _ce:
                        print("[OCR] Claude Vision error: " + str(_ce))

                # Если ни Tesseract, ни EasyOCR, ни Claude Vision не справились —
                # выдаём осмысленное сообщение пользователю.
                if not text.strip():
                    raise HTTPException(400,
                        "PDF содержит только сканы. Для распознавания установи один из:\n"
                        "• Tesseract OCR (запусти install_tesseract.bat)\n"
                        "• pip install easyocr\n"
                        "• добавь ANTHROPIC_API_KEY в .env для Claude Vision OCR")

        except HTTPException:
            raise
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(400, f"Ошибка чтения PDF: {e}")


    # ── DOCX / DOC ────────────────────────────────────────────────────────
    elif ext in ("docx", "doc"):
        try:
            import docx as _docx
            document   = _docx.Document(io.BytesIO(raw))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = "\n".join(paragraphs)
            print(f"[DOCX] {len(text)} символов")
        except Exception as e:
            raise HTTPException(400, f"Ошибка чтения DOCX: {e}")

    # ── TXT / MD / CSV / JSON ──────────────────────────────────────────────
    else:
        for enc in ("utf-8", "cp1251", "latin-1"):
            try:
                text = raw.decode(enc)
                break
            except (UnicodeDecodeError, AttributeError):
                continue

    if not text or not text.strip():
        raise HTTPException(
            400, "Не удалось извлечь текст. Поддерживаются: .txt .md .pdf .docx"
        )

    text = text.replace("\x00", "")
    name = filename.rsplit(".", 1)[0]
    doc  = await kb.add_document_async(name, text, category)

    print(f"[UPLOAD] OK {filename} -> {len(text)} символов")

    # ── НОВОЕ: уведомляем в Telegram о крупных документах (>500KB) ────
    if tg_on() and len(text) > 500_000:
        try:
            await tg_notify(
                f"📄 Загружен крупный документ\n"
                f"Имя: <b>{name}</b>\n"
                f"Категория: <code>{category}</code>\n"
                f"Размер: <b>{len(text):,}</b> символов",
                level="info", silent=True,
            )
        except Exception:
            pass

    # Auto-summary в фоне
    async def _bg_summary():
        try:
            await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(
                    None, lambda: orchestrator.synthesis.summarize_document(text[:2000], name)
                    if hasattr(orchestrator, 'synthesis') and hasattr(orchestrator.synthesis, 'summarize_document')
                    else None
                ), timeout=15.0
            )
        except Exception:
            pass
    asyncio.create_task(_bg_summary())

    return {
        "success":  True,
        "id":       doc.get("id"),
        "name":     name,
        "category": category,
        "size":     len(text),
        "chunks":   doc.get("chunk_count", doc.get("chunks", 0)),
        "message":  f"Документ «{name}» добавлен в базу знаний",
    }


@app.delete("/api/kb/documents/{doc_id}", tags=["knowledge"])
async def kb_delete(doc_id: int):
    await kb.delete_document_async(doc_id)
    return {"success": True, "stats": kb.stats()}


@app.post("/api/kb/chat", tags=["knowledge"])
async def kb_chat(req: ChatRequest):
    result = kb.answer(req.question.strip())
    await db_save_chat(
        session_id=req.session_id,
        question=req.question.strip(),
        answer=result["answer"],
        sources=result.get("sources", []),
        confidence=result.get("confidence", 0),
        found=result.get("found", False),
    )
    return result


@app.get("/api/kb/chat/history", tags=["knowledge"])
async def kb_chat_history(session_id: str = "default", limit: int = 50):
    history = await db_get_chat_history(session_id, limit)
    return {"history": history, "session_id": session_id}


@app.post("/api/kb/search", tags=["knowledge"])
def kb_search(req: ChatRequest):
    results = kb.search(req.question.strip(), top_k=5)
    return {"results": results, "total": len(results)}


@app.post("/api/kb/chat/fast", tags=["kb"])
async def kb_chat_fast(req: ChatRequest):
    """
    LEGACY endpoint — перенаправляет на /api/kb/chat/fast/v2 с MCP.
    Сохраняет совместимость с фронтом v4.0, но теперь использует:
      • Hybrid RAG (semantic + keyword)
      • MCP integration (legal_rk, hr, docs)
      • Multi-LLM fallback (Groq → Gemini → Anthropic → Ollama)
    """
    # Импортируем здесь чтобы избежать circular imports
    from chat_fast_v2 import ChatRequest as ChatRequestV2

    # Найдём v2 endpoint через app.routes
    v2_endpoint = None
    for route in app.routes:
        if getattr(route, "path", "") == "/api/kb/chat/fast/v2":
            v2_endpoint = route.endpoint
            break

    if v2_endpoint is None:
        logger.error("[chat_fast LEGACY] /api/kb/chat/fast/v2 not found in app.routes")
        return {
            "answer": "⚠️ Внутренняя ошибка: chat_fast_v2 не зарегистрирован.",
            "found": False, "sources": [], "confidence": 0,
            "mode": "error", "flags": ["v2_not_registered"],
        }

    # Конвертируем legacy request в v2 формат с включённым MCP
    v2_req = ChatRequestV2(
        question=req.question,
        session_id=req.session_id or "default",
        use_cache=True,
        rerank=False,   # выключено по умолчанию (модель 2.27GB)
        use_web=True,
        use_mcp=True,   # ВКЛЮЧАЕМ MCP — главное отличие
    )
    logger.info("[chat_fast LEGACY] redirect → v2 (MCP enabled)")
    return await v2_endpoint(v2_req)


@app.post("/api/agent/chat", tags=["agents"])
async def agent_chat(req: AgentChatRequest):
    """LEGACY endpoint → также редиректит на chat_fast_v2 с MCP."""
    fast_req = ChatRequest(question=req.question, session_id=req.session_id)
    return await kb_chat_fast(fast_req)


@app.get("/api/agent/graph", tags=["agents"])
async def agent_graph_schema():
    return orchestrator.get_graph_schema()


@app.get("/api/agent/metrics", tags=["agents"])
def agent_metrics():
    if orchestrator is None:
        return {"error": "Orchestrator not ready"}
    return orchestrator.metrics.summary()


@app.get("/api/agent/traces", tags=["agents"])
def agent_traces(limit: int = 20):
    if orchestrator is None:
        return {"traces": []}
    return {"traces": orchestrator.recent_traces(limit), "total": 0}


@app.post("/api/predict/ocr", tags=["vision"])
async def predict_ocr(file: UploadFile = File(...)):
    raw = await file.read()
    text = ""
    try:
        import numpy as np
        from PIL import Image as PILImage
        img = PILImage.open(io.BytesIO(raw)).convert("RGB")
        try:
            import easyocr
            reader  = easyocr.Reader(["ru", "en"], gpu=False, verbose=False)
            results = reader.readtext(np.array(img), detail=0, paragraph=True)
            text    = "\n".join(results)
        except ImportError:
            try:
                import pytesseract
                text = pytesseract.image_to_string(img, lang="rus+eng")
            except Exception:
                text = ""
    except Exception as e:
        logger.warning("OCR failed: %s", e)
    return {"text": text.strip(), "chars": len(text.strip())}


# /api/auth/send-code теперь в auth_routes.py (build_auth_router)


# ── НОВОЕ: системные эндпоинты для интеграций ─────────────────────────

@app.get("/api/system/cache", tags=["system"])
def system_cache():
    return cache_stats()


@app.get("/api/system/reranker", tags=["system"])
def system_reranker():
    return reranker_status()


@app.get("/api/system/integrations", tags=["system"])
def system_integrations():
    return {
        "cache":    cache_stats(),
        "reranker": reranker_status(),
        "telegram": {"enabled": tg_on()},
        "train_ws": {"subscribers": train_ws.count()},
        "auth":     {"endpoints": ["/api/auth/register", "/api/auth/login",
                                   "/api/auth/me", "/api/auth/logout"]},
    }